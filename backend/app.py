from flask import Flask, request, jsonify, render_template
import PyPDF2
from docx import Document
from dotenv import load_dotenv
import requests
import json
import re
import os
import uuid
from datetime import datetime

load_dotenv()

app = Flask(
    __name__,
    template_folder="../frontend",
    static_folder="../frontend",
    static_url_path=""
)

UPLOAD_FOLDER = "uploads"
HISTORY_FOLDER = "history"
HISTORY_FILE = os.path.join(HISTORY_FOLDER, "sets.json")

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(HISTORY_FOLDER, exist_ok=True)

ANSWER_START = "[[ANSWER]]"
ANSWER_END = "[[/ANSWER]]"

OLLAMA_URL = os.getenv("OLLAMA_URL", "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "openai/localmodel:latest")


@app.route("/")
def home():
    return render_template("index.html")


# -----------------------------
# BASIC HELPERS
# -----------------------------

def normalize_space(text):
    return re.sub(r"\s+", " ", str(text)).strip()


def safe_int(value, default=20, minimum=1, maximum=50):
    try:
        number = int(value)
        return max(minimum, min(number, maximum))
    except Exception:
        return default


def has_answer_marker(text):
    return (
        ANSWER_START in text
        or ANSWER_END in text
        or bool(re.search(r"<u>.*?</u>", text, flags=re.IGNORECASE))
        or bool(re.search(r"\*\*.*?\*\*", text))
    )


def remove_answer_markers(text):
    text = str(text).replace(ANSWER_START, "").replace(ANSWER_END, "")
    text = re.sub(r"</?u>", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    return normalize_space(text)


def clean_ai_suffix(text):
    return normalize_space(str(text).replace("(AI-generated answer)", ""))


# -----------------------------
# HISTORY STORAGE
# -----------------------------

def load_history():
    if not os.path.exists(HISTORY_FILE):
        return []

    try:
        with open(HISTORY_FILE, "r", encoding="utf-8") as file:
            return json.load(file)
    except Exception:
        return []


def save_history(history):
    with open(HISTORY_FILE, "w", encoding="utf-8") as file:
        json.dump(history, file, indent=2, ensure_ascii=False)


def add_history_set(filename, mode, cards):
    history = load_history()

    set_data = {
        "id": str(uuid.uuid4()),
        "filename": filename,
        "mode": mode,
        "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "count": len(cards),
        "cards": cards
    }

    history.insert(0, set_data)
    history = history[:30]

    save_history(history)
    return set_data


@app.route("/history", methods=["GET"])
def get_history():
    history = load_history()

    summaries = [
        {
            "id": item["id"],
            "filename": item["filename"],
            "mode": item["mode"],
            "created_at": item["created_at"],
            "count": item["count"]
        }
        for item in history
    ]

    return jsonify({"history": summaries})


@app.route("/history/<set_id>", methods=["GET"])
def get_history_set(set_id):
    history = load_history()

    for item in history:
        if item["id"] == set_id:
            return jsonify(item)

    return jsonify({"error": "History set not found"}), 404


@app.route("/history/<set_id>", methods=["PUT"])
def update_history_set(set_id):
    data = request.get_json(force=True)
    cards = data.get("cards", [])

    history = load_history()

    for item in history:
        if item["id"] == set_id:
            item["cards"] = cards
            item["count"] = len(cards)
            save_history(history)
            return jsonify({"success": True})

    return jsonify({"error": "History set not found"}), 404


@app.route("/history/<set_id>", methods=["DELETE"])
def delete_history_set(set_id):
    history = load_history()
    new_history = [item for item in history if item["id"] != set_id]

    save_history(new_history)
    return jsonify({"success": True})


# -----------------------------
# FILE EXTRACTION
# -----------------------------

def extract_text_from_pdf(file_path):
    text = ""

    with open(file_path, "rb") as file:
        reader = PyPDF2.PdfReader(file)

        for page_number, page in enumerate(reader.pages, start=1):
            extracted = page.extract_text()
            if extracted:
                text += f"\n[[PAGE {page_number}]]\n"
                text += extracted + "\n"

    return text


def paragraph_text_with_markers(paragraph):
    parts = []

    for run in paragraph.runs:
        if not run.text:
            continue

        if run.underline or run.font.highlight_color:
            parts.append(f"{ANSWER_START}{run.text}{ANSWER_END}")
        else:
            parts.append(run.text)

    return "".join(parts).strip()


def extract_text_from_docx(file_path):
    doc = Document(file_path)
    lines = []

    for para in doc.paragraphs:
        line = paragraph_text_with_markers(para)
        if line:
            lines.append(line)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    line = paragraph_text_with_markers(para)
                    if line:
                        lines.append(line)

    return "\n".join(lines)


# -----------------------------
# QUESTIONNAIRE DETECTION
# -----------------------------

def split_inline_questions_and_choices(text):
    text = text.replace("\r", "\n")

    text = re.sub(r"\s+(?=(?:#+\s*)?\d{1,3}[\.\)]\s+)", "\n", text)
    text = re.sub(r"\s+(?=\(?[a-dA-D][\)\.]\s+)", "\n", text)

    return text


def looks_like_quiz_question(question):
    q = question.lower().strip()

    clues = [
        "what",
        "which",
        "who",
        "when",
        "where",
        "why",
        "how",
        "it is",
        "it refers",
        "it means",
        "the following",
        "except",
        "classify",
        "identify",
        "choose",
        "select",
        "_____",
    ]

    if question.endswith("?"):
        return True

    return any(clue in q for clue in clues)


def parse_multiple_choice_questions(text):
    text = split_inline_questions_and_choices(text)

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    cards = []
    current = None

    question_pattern = re.compile(r"^(?:#+\s*)?(\d{1,3})[\.\)]\s*(.+)$")
    option_pattern = re.compile(r"^\(?([a-dA-D])[\)\.]\s*(.+)$")

    def save_current():
        nonlocal current

        if not current:
            return

        question = remove_answer_markers(current["question"])
        choices = [remove_answer_markers(choice) for choice in current["choices"]]
        answer = remove_answer_markers(current["answer"])

        if (
            question
            and looks_like_quiz_question(question)
            and 2 <= len(choices) <= 6
            and len(question) <= 500
        ):
            cards.append({
                "question": question,
                "choices": choices,
                "answer": answer if answer else "No answer marked in the file.",
                "answer_source": "file" if answer else "missing"
            })

        current = None

    for line in lines:
        if line == "---":
            continue

        question_match = question_pattern.match(line)

        if question_match:
            save_current()

            current = {
                "number": question_match.group(1),
                "question": question_match.group(2),
                "choices": [],
                "answer": ""
            }
            continue

        option_match = option_pattern.match(line)

        if option_match and current:
            letter = option_match.group(1).lower()
            option_text = option_match.group(2)

            full_choice = f"({letter}) {option_text}"
            current["choices"].append(full_choice)

            if has_answer_marker(option_text):
                current["answer"] = full_choice

            continue

        if current:
            if current["choices"]:
                current["choices"][-1] += " " + line

                if has_answer_marker(line):
                    current["answer"] = current["choices"][-1]
            else:
                current["question"] += " " + line

    save_current()
    return cards


def should_use_questionnaire_mode(cards):
    if not cards:
        return False

    marked_count = sum(
        1 for card in cards
        if card.get("answer_source") == "file"
    )

    if marked_count >= 2:
        return True

    if len(cards) >= 3:
        return True

    return False


# -----------------------------
# LOCAL AI: ANSWER MCQ
# -----------------------------

def extract_json_from_ai_response(response_text):
    response_text = str(response_text).strip()

    try:
        return json.loads(response_text)
    except json.JSONDecodeError:
        pass

    object_match = re.search(r"{.*}", response_text, flags=re.DOTALL)
    if object_match:
        try:
            return json.loads(object_match.group(0))
        except json.JSONDecodeError:
            pass

    array_match = re.search(r"\[\s*{.*?}\s*\]", response_text, flags=re.DOTALL)
    if array_match:
        try:
            return json.loads(array_match.group(0))
        except json.JSONDecodeError:
            pass

    return None


def ask_ollama_to_answer_mcq(question, choices):
    choices_text = "\n".join(choices)

    prompt = f"""
You are answering a multiple-choice quiz.

Choose the single best answer from the choices.

Return ONLY valid JSON:
{{
  "answer_letter": "a",
  "answer": "(a) full answer text here"
}}

Rules:
- Choose only one option.
- Copy the answer exactly from the choices.
- Do not explain.
- Do not use markdown.

Question:
{question}

Choices:
{choices_text}
"""

    body = {
        "model": OLLAMA_MODEL,
        "prompt": prompt,
        "stream": False,
        "format": "json",
        "options": {
            "temperature": 0.05
        }
    }

    response = requests.post(
        f"{OLLAMA_URL}/api/generate",
        json=body,
        timeout=180
    )

    response.raise_for_status()

    result = response.json()
    ai_text = result.get("response", "")
    data = extract_json_from_ai_response(ai_text)

    if not data:
        return "AI could not determine the answer."

    ai_answer = normalize_space(data.get("answer", ""))
    ai_letter = normalize_space(data.get("answer_letter", "")).lower()

    for choice in choices:
        if ai_answer.lower() == normalize_space(choice).lower():
            return normalize_space(choice)

    if ai_letter:
        for choice in choices:
            if choice.lower().startswith(f"({ai_letter})"):
                return normalize_space(choice)

    return "AI could not determine the answer."


def fill_missing_questionnaire_answers(cards):
    updated_cards = []

    for card in cards:
        if card.get("answer_source") == "file":
            updated_cards.append(card)
            continue

        try:
            ai_answer = ask_ollama_to_answer_mcq(
                card["question"],
                card["choices"]
            )

            card["answer"] = ai_answer
            card["answer_source"] = "ai"
        except Exception as error:
            print("AI answer error:", error)
            card["answer"] = "AI could not determine the answer."
            card["answer_source"] = "ai"

        updated_cards.append(card)

    return updated_cards


# -----------------------------
# LESSON TEXT CLEANING
# -----------------------------

def clean_lesson_line(line):
    line = normalize_space(line)

    line = re.sub(r"\[\[PAGE\s+\d+\]\]", "", line)
    line = re.sub(r"^\d+\s*\|\s*Page.*$", "", line, flags=re.IGNORECASE)
    line = re.sub(r"HCDC[-\w.]*", "", line, flags=re.IGNORECASE)
    line = re.sub(r"</?u>", "", line, flags=re.IGNORECASE)

    return normalize_space(line)


def is_bad_lesson_line(line):
    if not line:
        return True

    lower = line.lower()

    bad_exact = {
        "introduction",
        "activity",
        "analysis",
        "application",
        "closure",
        "objectives",
        "references",
        "instructions",
        "materials needed",
        "course outline",
        "course information",
        "grading system",
        "guidelines",
        "table of contents",
        "module overview",
    }

    if lower in bad_exact:
        return True

    bad_phrases = [
        "course pack in",
        "this is a property",
        "no part of this course pack",
        "reproduced or photocopied",
        "authorized school administrators",
        "at the end of the lesson",
        "at the end of the module",
        "students will create",
        "split your class",
        "ask each group",
        "provide them",
        "guide questions",
        "the following questions",
        "well done",
        "you’re now ready",
        "feel free to continue",
        "prepared by",
        "reviewed by",
        "approved by",
        "endorsed by",
        "jornie joy",
        "diosdado",
        "holy cross of davao college",
    ]

    if any(phrase in lower for phrase in bad_phrases):
        return True

    if len(line) < 40:
        return True

    return False


def extract_study_text(text):
    raw_lines = text.splitlines()
    lines = []

    for line in raw_lines:
        cleaned = clean_lesson_line(line)
        if cleaned:
            lines.append(cleaned)

    study_lines = []
    collecting = False

    start_headers = {"abstraction"}
    stop_headers = {
        "application",
        "closure",
        "activity",
        "analysis",
        "introduction",
        "objectives",
        "references",
    }

    for line in lines:
        lower = line.lower().strip()

        if lower in start_headers:
            collecting = True
            continue

        if lower in stop_headers:
            collecting = False
            continue

        if collecting and not is_bad_lesson_line(line):
            study_lines.append(line)

    study_text = " ".join(study_lines)
    study_text = normalize_space(study_text)

    if len(study_text) < 1200:
        fallback_lines = []

        for line in lines:
            if not is_bad_lesson_line(line):
                fallback_lines.append(line)

        study_text = " ".join(fallback_lines)
        study_text = normalize_space(study_text)

    study_text = re.sub(r"\([a-dA-D]\)\s+[^.?!]{1,120}", " ", study_text)
    study_text = re.sub(r"\s+", " ", study_text)

    return study_text.strip()


def chunk_text(text, max_chars=3500):
    words = text.split()
    chunks = []
    current = []
    current_length = 0

    for word in words:
        current.append(word)
        current_length += len(word) + 1

        if current_length >= max_chars:
            chunks.append(" ".join(current))
            current = []
            current_length = 0

    if current:
        chunks.append(" ".join(current))

    return chunks


# -----------------------------
# LOCAL AI: GENERATE FLASHCARDS
# -----------------------------

def normalize_ai_flashcards(data):
    if isinstance(data, dict):
        if "flashcards" in data and isinstance(data["flashcards"], list):
            data = data["flashcards"]
        elif "cards" in data and isinstance(data["cards"], list):
            data = data["cards"]
        elif "question" in data and "answer" in data:
            data = [data]
        else:
            return []

    if not isinstance(data, list):
        return []

    cards = []

    for item in data:
        if not isinstance(item, dict):
            continue

        question = normalize_space(item.get("question", ""))
        answer = normalize_space(item.get("answer", ""))

        if not question or not answer:
            continue

        if "main idea of this passage" in question.lower():
            continue

        if len(question) < 10 or len(answer) < 30:
            continue

        cards.append({
            "question": question,
            "choices": [],
            "answer": answer,
            "answer_source": "ai"
        })

    return cards


def ask_ollama_for_flashcards(text, count=8):
    prompt = f"""
You are creating study flashcards from a lesson/course pack.

Return ONLY valid JSON.

Return this exact JSON structure:
{{
  "flashcards": [
    {{
      "question": "clear study question here",
      "answer": "short but complete answer here"
    }}
  ]
}}

Rules:
- Create exactly {count} flashcards if possible.
- Do NOT use the question "What is the main idea of this passage?"
- Do NOT copy page headers, course pack title, author names, table of contents, objectives, activities, instructions, or references.
- Do NOT create questions from course information, grading system, or file metadata.
- Focus only on real lesson concepts, definitions, explanations, and important ideas.
- Questions must sound natural, like a teacher made them.
- Answers must be based only on the provided lesson text.
- Keep each answer between 1 and 3 sentences.
- Do not include markdown.

Lesson text:
{text}
"""

    body = {
        "model": OLLAMA_MODEL,
        "prompt": prompt,
        "stream": False,
        "format": "json",
        "options": {
            "temperature": 0.1
        }
    }

    response = requests.post(
        f"{OLLAMA_URL}/api/generate",
        json=body,
        timeout=300
    )

    response.raise_for_status()

    result = response.json()
    ai_text = result.get("response", "")

    data = extract_json_from_ai_response(ai_text)
    return normalize_ai_flashcards(data)


def generate_ai_flashcards_from_lesson(text, requested_count):
    study_text = extract_study_text(text)

    print("STUDY TEXT LENGTH:", len(study_text))
    print("STUDY TEXT PREVIEW:", study_text[:800])

    if len(study_text) < 500:
        return []

    chunks = chunk_text(study_text, max_chars=3500)

    all_cards = []
    used_questions = set()

    for chunk in chunks[:5]:
        needed = max(3, min(8, requested_count - len(all_cards)))

        if needed <= 0:
            break

        try:
            cards = ask_ollama_for_flashcards(chunk, count=needed)
        except Exception as error:
            print("OLLAMA FLASHCARD ERROR:", error)
            continue

        for card in cards:
            q_key = card["question"].lower()

            if q_key in used_questions:
                continue

            all_cards.append(card)
            used_questions.add(q_key)

            if len(all_cards) >= requested_count:
                return all_cards

    return all_cards


# -----------------------------
# MAIN GENERATION
# -----------------------------

def generate_questions(text, requested_count):
    questionnaire_cards = parse_multiple_choice_questions(text)

    if should_use_questionnaire_mode(questionnaire_cards):
        cards = questionnaire_cards[:requested_count]
        return "questionnaire", fill_missing_questionnaire_answers(cards)

    ai_cards = generate_ai_flashcards_from_lesson(text, requested_count)

    if ai_cards:
        return "lesson", ai_cards

    return "lesson", []


@app.route("/process", methods=["POST"])
def process_file():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    requested_count = safe_int(request.form.get("count", 20), default=20, maximum=50)

    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    file_path = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(file_path)

    if file.filename.lower().endswith(".pdf"):
        text = extract_text_from_pdf(file_path)
    elif file.filename.lower().endswith(".docx"):
        text = extract_text_from_docx(file_path)
    else:
        return jsonify({"error": "Unsupported file format. Please upload PDF or DOCX."}), 400

    mode, cards = generate_questions(text, requested_count)

    if not cards:
        return jsonify({
            "error": "No usable questions found. Make sure Ollama is running and the uploaded file has readable text."
        }), 500

    set_data = add_history_set(file.filename, mode, cards)

    return jsonify({
        "success": True,
        "set_id": set_data["id"],
        "filename": file.filename,
        "mode": mode,
        "questions": cards
    })


if __name__ == "__main__":
    app.run(debug=True)